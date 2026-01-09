"""Export utilities for transcripts (PDF, DOCX)."""

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Tuple

# Lazy imports for export libraries
def _get_docx():
    """Lazy import for python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        return Document, Inches, Pt, WD_ALIGN_PARAGRAPH, RGBColor
    except ImportError:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

def _get_reportlab():
    """Lazy import for reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
        return A4, getSampleStyleSheet, ParagraphStyle, inch, SimpleDocTemplate, Paragraph, Spacer, TA_LEFT, TA_JUSTIFY, ListFlowable, ListItem
    except ImportError:
        raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")


def _parse_markdown_line(line: str) -> Tuple[str, dict]:
    """Parse a single line of markdown and return (text, metadata).
    
    Returns:
        Tuple of (clean_text, metadata_dict) where metadata contains:
        - heading_level: 1-6 for headings, 0 for normal
        - is_bullet: True if list item
        - is_bold: True if entire line is bold
    """
    metadata = {"heading_level": 0, "is_bullet": False, "is_bold": False}
    text = line.strip()
    
    # Check for headings (# to ######)
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', text)
    if heading_match:
        metadata["heading_level"] = len(heading_match.group(1))
        text = heading_match.group(2)
        return text, metadata
    
    # Check for bullet points
    bullet_match = re.match(r'^[-*+]\s+(.+)$', text)
    if bullet_match:
        metadata["is_bullet"] = True
        text = bullet_match.group(1)
    
    # Check for numbered list
    numbered_match = re.match(r'^\d+\.\s+(.+)$', text)
    if numbered_match:
        metadata["is_bullet"] = True
        text = numbered_match.group(1)
    
    return text, metadata


def _apply_inline_formatting_docx(paragraph, text):
    """Apply inline markdown formatting (**bold**, *italic*, etc.) to a docx paragraph."""
    from docx.shared import Pt
    
    # Pattern to match **bold**, *italic*, ***bold-italic***
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Determine formatting
        if match.group(2):  # ***bold-italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # `code`
            run = paragraph.add_run(match.group(5))
            run.font.name = 'Consolas'
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _convert_markdown_to_html(text: str) -> str:
    """Convert basic markdown to HTML for reportlab."""
    # Escape HTML first
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    
    # Bold-italic ***text***
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
    # Bold **text**
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    # Italic *text*
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Code `text`
    text = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', text)
    
    return text


def export_to_docx(
    title: str,
    content: str,
    summary: Optional[str] = None,
    date: Optional[str] = None,
    duration: Optional[str] = None,
) -> bytes:
    """Export transcript to DOCX format with markdown support.
    
    Args:
        title: Document title
        content: Transcript text content
        summary: Optional summary text (markdown supported)
        date: Recording date
        duration: Recording duration
        
    Returns:
        DOCX file as bytes
    """
    Document, Inches, Pt, WD_ALIGN_PARAGRAPH, RGBColor = _get_docx()
    
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    if date or duration:
        meta_text = []
        if date:
            meta_text.append(f"Date: {date}")
        if duration:
            meta_text.append(f"Duration: {duration}")
        meta_para = doc.add_paragraph(" | ".join(meta_text))
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in meta_para.runs:
            run.font.size = Pt(10)
            run.font.italic = True
    
    doc.add_paragraph()  # Spacer
    
    # Summary section with markdown parsing
    if summary:
        doc.add_heading("Summary", level=1)
        
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                continue
                
            text, metadata = _parse_markdown_line(line)
            
            if metadata["heading_level"] > 0:
                # Add as heading (level 2-4 in summary)
                level = min(metadata["heading_level"] + 1, 4)
                doc.add_heading(text, level=level)
            elif metadata["is_bullet"]:
                # Add as bullet point
                p = doc.add_paragraph(style='List Bullet')
                _apply_inline_formatting_docx(p, text)
            else:
                # Regular paragraph with inline formatting
                p = doc.add_paragraph()
                _apply_inline_formatting_docx(p, text)
                p.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # Spacer
    
    # Transcript section
    doc.add_heading("Transcript", level=1)
    
    # Parse speaker labels and format
    for para in content.split("\n\n"):
        if para.strip():
            # Check for speaker label
            if para.strip().startswith("[Speaker"):
                # Extract speaker and text
                parts = para.split(":", 1)
                if len(parts) == 2:
                    speaker = parts[0].strip()
                    text = parts[1].strip()
                    p = doc.add_paragraph()
                    run = p.add_run(f"{speaker}: ")
                    run.bold = True
                    p.add_run(text)
                else:
                    doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph(para.strip())
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_to_pdf(
    title: str,
    content: str,
    summary: Optional[str] = None,
    date: Optional[str] = None,
    duration: Optional[str] = None,
) -> bytes:
    """Export transcript to PDF format with markdown support.
    
    Args:
        title: Document title
        content: Transcript text content
        summary: Optional summary text (markdown supported)
        date: Recording date
        duration: Recording duration
        
    Returns:
        PDF file as bytes
    """
    A4, getSampleStyleSheet, ParagraphStyle, inch, SimpleDocTemplate, Paragraph, Spacer, TA_LEFT, TA_JUSTIFY, ListFlowable, ListItem = _get_reportlab()
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=18,
        spaceAfter=12,
    )
    
    heading1_style = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
    )
    
    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontSize=11,
        spaceBefore=10,
        spaceAfter=4,
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        leftIndent=20,
        bulletIndent=10,
        spaceAfter=4,
    )
    
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        textColor='grey',
        alignment=1,  # Center
        spaceAfter=20,
    )
    
    speaker_style = ParagraphStyle(
        'Speaker',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    
    story = []
    
    # Title
    story.append(Paragraph(title, title_style))
    
    # Metadata
    if date or duration:
        meta_parts = []
        if date:
            meta_parts.append(date)
        if duration:
            meta_parts.append(duration)
        story.append(Paragraph(" | ".join(meta_parts), meta_style))
    
    # Summary with markdown parsing
    if summary:
        story.append(Paragraph("Summary", heading1_style))
        
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            
            text, metadata = _parse_markdown_line(line)
            html_text = _convert_markdown_to_html(text)
            
            if metadata["heading_level"] == 1:
                story.append(Paragraph(html_text, heading2_style))
            elif metadata["heading_level"] == 2:
                story.append(Paragraph(html_text, heading2_style))
            elif metadata["heading_level"] >= 3:
                story.append(Paragraph(html_text, heading3_style))
            elif metadata["is_bullet"]:
                story.append(Paragraph(f"• {html_text}", bullet_style))
            else:
                story.append(Paragraph(html_text, body_style))
        
        story.append(Spacer(1, 12))
    
    # Transcript
    story.append(Paragraph("Transcript", heading1_style))
    
    for para in content.split("\n\n"):
        if para.strip():
            # Check for speaker label
            if para.strip().startswith("[Speaker"):
                parts = para.split(":", 1)
                if len(parts) == 2:
                    speaker = parts[0].strip().replace("[", "").replace("]", "")
                    text = parts[1].strip()
                    # Escape HTML
                    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    formatted = f"<b>[{speaker}]:</b> {text}"
                    story.append(Paragraph(formatted, speaker_style))
                else:
                    safe_para = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe_para, body_style))
            else:
                safe_para = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_para, body_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.read()
